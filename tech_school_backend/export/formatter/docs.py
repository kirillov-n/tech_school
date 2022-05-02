from export.formatter.base import AbstractFormatter

from docx.oxml.xmlchemy import OxmlElement
from docx.shared import RGBColor
from docx.oxml.shared import qn
from docx.table import Table
import docx


class DOCSFormatter(AbstractFormatter):
    def __init__(self, *args, **kwargs):
        super(DOCSFormatter, self).__init__(*args, **kwargs)
        self.document = docx.Document()

    def add_date(self):
        self.document.add_paragraph().add_run(str(self.date)).bold = True

    def add_table(self):
        columns = len(self.fields)
        rows = len(self.queryset) + 2

        table = self.document.add_table(rows=rows, cols=columns)
        table.style = "Table Grid"

        super(DOCSFormatter, self).add_table(table)

    def add_top_fields(self, table: Table):
        col = 0
        for field in self.verbose_fields:
            cell = table.cell(0, col)
            self._set_cell_background(cell, "708090")
            run = cell.add_paragraph().add_run(field)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            col += 1

    def add_all_fields(self, table: Table):
        row = 1
        for obj in self.queryset:
            column = 0
            for field in self.fields:
                cell = table.cell(row, column)
                cell.text = str(getattr(obj, field, ""))
                column += 1
            row += 1

    def save(self):
        self.document.save(self.buffer)

    @staticmethod
    def _set_cell_background(cell, fill: str):
        cell_properties = cell._element.tcPr
        try:
            cell_shading = cell_properties.xpath('w:shd')[0]
        except IndexError:
            cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), fill)
        cell_properties.append(cell_shading)
